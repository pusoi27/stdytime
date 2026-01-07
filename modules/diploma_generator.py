"""
Diploma/Certificate Generation Module
Integrates diploma generation from DOCX templates and PDF conversion

This module provides utilities for:
- Generating diplomas from DOCX templates with placeholder replacement
- Converting DOCX certificates to PDF
- Managing multiple diploma types (Award, Recognition, Welcome)
"""

from pathlib import Path
from datetime import datetime
from typing import Dict, Any, Optional, List, Set
import re

try:
    import pandas as pd
    HAS_PANDAS = True
except ImportError:
    HAS_PANDAS = False

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    from docx2pdf import convert
    HAS_DOCX2PDF = True
except ImportError:
    HAS_DOCX2PDF = False

try:
    import win32com.client as win32
    HAS_WIN32COM = True
except ImportError:
    HAS_WIN32COM = False


# Template mapping for diploma types
TEMPLATE_MAP = {
    "Award": "Certificate of Award.docx",
    "Certificate": "Certificate of Recognition.docx",
    "Welcome": "Certificate of Welcome.docx",
}

# Placeholder tokens in templates
PLACEHOLDERS = {
    "name": "[[NAME]]",
    "diploma": "[[DIPLOMA]]",
    "date": "[[DATE]]",
    "subjects": "[[SUBJECTS]]",
    "success": "[[SUCCESS]]",
}


def _letter_part(level: Optional[str]) -> str:
    """Extract the base level letter part (e.g., 'F70' -> 'F', 'AI20' -> 'AI')."""
    if not level:
        return ""
    s = str(level).strip().upper().replace(' ', '')
    
    # Use regex to extract letter+digit portion before page_index
    m = re.match(r'^([A-Z0-9]+?)(\d+)$', s)
    if m:
        return m.group(1)
    return s


def _completed_level(current_level: Optional[str]) -> str:
    """Get the completed level from current level (e.g., 'F70' -> 'E', '3A20' -> '4A')."""
    if not current_level:
        return ""
    
    base = _letter_part(current_level)
    if not base:
        return ""
    
    # Reading/Math level progression
    # For numbered levels (7A-2A), decrement goes backward (3A -> 4A completed)
    # For letter levels (A-L), go back one letter (F -> E completed)
    
    # Check if it's a numbered level like 7A, 6A, etc.
    if re.match(r'^\d+A$', base):
        num = int(base[:-1])
        if num < 7:
            return f"{num + 1}A"
        return base  # 7A has no previous
    
    # Single letter levels A-L
    if len(base) == 1 and base.isalpha():
        if base == 'A':
            return '2A'
        # Go back one letter: B->A, C->B, F->E, etc.
        prev_letter = chr(ord(base) - 1)
        return prev_letter
    
    # Double letter levels (AI, AII, BI, etc.)
    if len(base) == 2 and base[0].isalpha() and base[1] in ('I', 'V'):
        if base.endswith('I'):
            # AI -> A (same letter, remove subdivisions)
            return base[0]
        # II -> AI, III -> AII, etc.
        return base[0] + 'I'
    
    # Multi-letter or special cases - return as-is
    return base


def _replace_placeholders(doc: 'Document', mapping: Dict[str, str]) -> None:
    """Replace placeholders in all paragraphs and table cells."""
    if not HAS_DOCX:
        return
    
    # Replace in paragraphs
    for p in doc.paragraphs:
        for key, value in mapping.items():
            token = PLACEHOLDERS.get(key)
            if token and token in p.text:
                # Replace across runs to preserve formatting
                for r in p.runs:
                    r.text = r.text.replace(token, value)
    
    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in mapping.items():
                    token = PLACEHOLDERS.get(key)
                    if token and token in cell.text:
                        for p in cell.paragraphs:
                            for r in p.runs:
                                r.text = r.text.replace(token, value)


def generate_diplomas(
    classified_csv: str,
    template_dir: str,
    output_dir: str,
    students_filter: Optional[List[str]] = None,
) -> List[Dict[str, str]]:
    """
    Generate diplomas from DOCX templates for students.
    
    Args:
        classified_csv: Path to CSV with columns: Full Name, Diploma, Subject, NormalizedLevel
        template_dir: Directory containing DOCX templates
        output_dir: Directory to save generated DOCX files
        students_filter: Optional list of student names to filter
    
    Returns:
        List of dictionaries with 'Full Name', 'Diploma', and 'Certificate' (file path)
    """
    if not HAS_DOCX or not HAS_PANDAS:
        return []
    
    csv_path = Path(classified_csv)
    tmpl_dir = Path(template_dir)
    out_dir = Path(output_dir)
    out_dir.mkdir(parents=True, exist_ok=True)
    
    # Read CSV
    try:
        df = pd.read_csv(csv_path)
    except Exception as e:
        print(f"Error reading CSV {csv_path}: {e}")
        return []
    
    # Filter by students if provided
    if students_filter:
        names_set: Set[str] = {n.strip() for n in students_filter}
        df = df[df['Full Name'].isin(names_set)]
    
    # Aggregate subjects per student
    subjects_map: Dict[str, str] = {}
    for name, group in df.groupby("Full Name"):
        subs = sorted(set(str(s).strip() for s in group.get("Subject", []) if pd.notna(s)))
        subjects_map[name] = ", ".join(subs) if subs else ""
    
    outputs: List[Dict[str, str]] = []
    today_str = datetime.now().strftime("%b %d, %Y")
    
    for name, group in df.groupby("Full Name"):
        # Get diploma type
        diploma = str(group["Diploma"].iloc[0]).strip() if "Diploma" in group.columns and len(group) > 0 else ""
        
        if not diploma:
            continue  # Skip students without diploma assignment
        
        tmpl_name = TEMPLATE_MAP.get(diploma)
        if not tmpl_name:
            continue  # Unknown diploma type
        
        tmpl_path = tmpl_dir / tmpl_name
        
        # Load or create template
        doc: Optional['Document'] = None
        if tmpl_path.exists():
            try:
                doc = Document(str(tmpl_path))
            except Exception:
                pass
        
        if doc is None:
            # Fallback: create simple document
            doc = Document()
            doc.add_heading('Certificate', level=0)
        
        # Extract subject levels for success text
        math_level = None
        reading_level = None
        math_current = None
        reading_current = None
        
        for _, row in group.iterrows():
            subj = str(row.get('Subject', '')).strip().lower()
            norm_level = row.get('NormalizedLevel') or row.get('Highest WS Completed This Month')
            
            completed = _completed_level(norm_level) if norm_level else ""
            current = _letter_part(norm_level) if norm_level else ""
            
            if subj.startswith('math'):
                math_level = completed
                math_current = current
            elif subj.startswith('reading'):
                reading_level = completed
                reading_current = current
        
        # Build success text based on diploma type
        subjects_text = subjects_map.get(name, "")
        
        if diploma == 'Welcome':
            # Welcome: show current level
            parts = []
            if math_current:
                parts.append(f"Kumon Math {math_current}")
            if reading_current:
                if parts:
                    parts.append(f"and Reading {reading_current}")
                else:
                    parts.append(f"Reading {reading_current}")
            success_text = " ".join(parts) + " Program" if parts else subjects_text
        else:
            # Award/Certificate: show completed level
            parts = []
            if math_level:
                parts.append(f"Kumon Math Level {math_level}")
            if reading_level:
                if parts:
                    parts.append(f"and Kumon Reading Level {reading_level}")
                else:
                    parts.append(f"Kumon Reading Level {reading_level}")
            success_text = " ".join(parts) if parts else subjects_text
        
        # Replace placeholders
        mapping = {
            "name": name,
            "diploma": diploma,
            "date": today_str,
            "subjects": subjects_text,
            "success": success_text,
        }
        
        try:
            _replace_placeholders(doc, mapping)
        except Exception as e:
            print(f"Warning: Placeholder replacement failed for {name}: {e}")
            # Fallback: add text directly
            doc.add_paragraph(f"Name: {name}")
            doc.add_paragraph(f"Diploma: {diploma}")
            doc.add_paragraph(f"Date: {today_str}")
            if subjects_text:
                doc.add_paragraph(f"Subjects: {subjects_text}")
        
        # Save output
        safe_name = name.replace("/", "-").replace("\\", "-")
        out_path = out_dir / f"{safe_name} - {diploma}.docx"
        
        try:
            doc.save(str(out_path))
        except PermissionError:
            # File locked, use unique filename
            unique = out_dir / f"{safe_name} - {diploma} - {datetime.now().strftime('%Y%m%d%H%M%S')}.docx"
            doc.save(str(unique))
            out_path = unique
        
        # Windows-only: Use Word COM to replace text in shapes/textboxes
        if HAS_WIN32COM:
            try:
                word = win32.Dispatch('Word.Application')
                word.Visible = False
                docx = word.Documents.Open(str(out_path.absolute()))
                
                for key, val in mapping.items():
                    token = PLACEHOLDERS.get(key)
                    targets = [token] if token else []
                    
                    # Add literal fallbacks
                    if key == 'name':
                        targets.append('JUNIOR PIERRE')
                    if key == 'success':
                        targets.extend([
                            'The Kumon Math Level 3A and Kumon Reading Levels 7A & 6A',
                            'Kumon Math Level 3A and Kumon Reading Levels 6A',
                            'Kumon subject Program'
                        ])
                    
                    for target in targets:
                        if not target:
                            continue
                        find = docx.Content.Find
                        find.Text = target
                        find.Replacement.Text = val
                        find.Forward = True
                        find.Wrap = 1  # wdFindContinue
                        find.Format = False
                        find.Execute(Replace=2)  # wdReplaceAll
                
                # Replace in shapes/text boxes
                for shp in docx.Shapes:
                    try:
                        if shp.TextFrame.HasText:
                            rng = shp.TextFrame.TextRange
                            text = rng.Text
                            for key, val in mapping.items():
                                token = PLACEHOLDERS.get(key)
                                if token and token in text:
                                    rng.Text = text.replace(token, val)
                                    text = rng.Text
                    except Exception:
                        continue
                
                docx.Save()
                docx.Close(False)
                word.Quit()
            except Exception as e:
                print(f"Warning: COM replacement failed: {e}")
        
        outputs.append({
            "Full Name": name,
            "Diploma": diploma,
            "Certificate": str(out_path)
        })
    
    return outputs


def convert_diplomas_to_pdf(
    diplomas: List[Dict[str, str]],
    output_dir: str,
) -> Dict[str, Any]:
    """
    Convert generated DOCX diplomas to PDF.
    
    Args:
        diplomas: List of diploma dicts with 'Certificate' file paths
        output_dir: Directory to save PDF files
    
    Returns:
        Dictionary with 'success_count', 'failed', and 'output_dir'
    """
    if not HAS_DOCX2PDF:
        return {
            'success_count': 0,
            'failed': [d.get('Full Name', 'Unknown') for d in diplomas],
            'error': 'docx2pdf not installed'
        }
    
    out_dir = Path(output_dir)
    out_dir.mkdir(parents=True, exist_ok=True)
    
    success_count = 0
    failed = []
    
    for diploma in diplomas:
        docx_path = Path(diploma.get('Certificate', ''))
        student_name = diploma.get('Full Name', 'Unknown')
        
        if not docx_path.exists():
            failed.append(student_name)
            continue
        
        pdf_path = out_dir / (docx_path.stem + '.pdf')
        
        try:
            convert(str(docx_path), str(pdf_path))
            success_count += 1
        except Exception as e:
            failed.append(f"{student_name}: {str(e)}")
    
    return {
        'success_count': success_count,
        'failed': failed,
        'output_dir': str(out_dir),
        'total': len(diplomas)
    }


def generate_and_convert_diplomas(
    classified_csv: str,
    template_dir: str,
    output_docx_dir: str,
    output_pdf_dir: str,
    students_filter: Optional[List[str]] = None,
) -> Dict[str, Any]:
    """
    Generate diplomas and convert to PDF in one call.
    
    Args:
        classified_csv: Path to classified students CSV
        template_dir: Directory with DOCX templates
        output_docx_dir: Directory for DOCX files
        output_pdf_dir: Directory for PDF files
        students_filter: Optional list of student names
    
    Returns:
        Summary dictionary with generation and conversion results
    """
    # Generate diplomas
    diplomas = generate_diplomas(
        classified_csv,
        template_dir,
        output_docx_dir,
        students_filter
    )
    
    if not diplomas:
        return {
            'status': 'error',
            'message': 'No diplomas generated',
            'generated_count': 0,
            'pdf_count': 0
        }
    
    # Convert to PDF
    pdf_result = convert_diplomas_to_pdf(diplomas, output_pdf_dir)
    
    return {
        'status': 'success' if pdf_result['success_count'] == len(diplomas) else 'partial',
        'generated_count': len(diplomas),
        'pdf_count': pdf_result['success_count'],
        'failed': pdf_result.get('failed', []),
        'docx_dir': output_docx_dir,
        'pdf_dir': pdf_result.get('output_dir', output_pdf_dir)
    }


# Export main functions
__all__ = [
    'generate_diplomas',
    'convert_diplomas_to_pdf',
    'generate_and_convert_diplomas'
]
